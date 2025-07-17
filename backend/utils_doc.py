from docx import Document
from docx.oxml.exceptions import InvalidXmlError
import regex as re
import logging
from typing import List, Dict, Any, Optional, Union
from io import BytesIO
import os
import io

# PDF处理库
try:
    import PyPDF2
    from PyPDF2.errors import PdfReadError
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyPDF2库未安装，PDF功能将不可用。请使用 'pip install PyPDF2' 安装。")

logger = logging.getLogger(__name__)

# 章节识别正则表达式
CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百千万]+章\s*[：:]?\s*(.*)$", re.MULTILINE),
    re.compile(r"^\d{1,2}\.\s*(.*)$", re.MULTILINE),
    re.compile(r"^[一二三四五六七八九十]+[、.]\s*(.*)$", re.MULTILINE),
    re.compile(r"^\([一二三四五六七八九十]+\)\s*(.*)$", re.MULTILINE),
    re.compile(r"^[A-Z]\s*[、.]\s*(.*)$", re.MULTILINE)
]

# 重要关键词模式
IMPORTANT_KEYWORDS = [
    "投标", "招标", "评标", "中标", "废标", "资格", "技术要求", "商务要求",
    "评分标准", "加分", "减分", "否决", "资质", "业绩", "注册资本", "人员",
    "设备", "品牌", "型号", "厂商", "供应商", "代理", "授权", "证书"
]

class DocumentProcessError(Exception):
    """文档处理异常"""
    pass

def validate_docx_file(file_obj) -> None:
    """验证docx文件格式"""
    try:
        # 重置文件指针
        file_obj.seek(0)
        
        # 尝试读取文件头
        header = file_obj.read(4)
        file_obj.seek(0)
        
        # 检查是否为ZIP格式（docx本质上是ZIP文件）
        if header != b'PK\x03\x04':
            raise DocumentProcessError("文件不是有效的docx格式")
            
    except Exception as e:
        raise DocumentProcessError(f"文件格式验证失败: {str(e)}")

def validate_pdf_file(file_obj) -> None:
    """验证PDF文件格式"""
    if not PDF_SUPPORT:
        raise DocumentProcessError("系统未安装PDF处理库，无法处理PDF文件")
        
    try:
        # 重置文件指针
        file_obj.seek(0)
        
        # 尝试读取文件头
        header = file_obj.read(5)
        file_obj.seek(0)
        
        # 检查是否为PDF格式（PDF文件以%PDF-开头）
        if header != b'%PDF-':
            raise DocumentProcessError("文件不是有效的PDF格式")
            
        # 尝试使用PyPDF2读取文件，验证结构
        try:
            reader = PyPDF2.PdfReader(file_obj)
            if len(reader.pages) == 0:
                raise DocumentProcessError("PDF文件不包含任何页面")
            file_obj.seek(0)  # 重置文件指针
        except PdfReadError as e:
            raise DocumentProcessError(f"PDF文件结构无效: {str(e)}")
            
    except Exception as e:
        raise DocumentProcessError(f"PDF文件格式验证失败: {str(e)}")

def extract_text_from_pdf(file_obj) -> str:
    """从PDF文件中提取文本"""
    if not PDF_SUPPORT:
        raise DocumentProcessError("系统未安装PDF处理库，无法处理PDF文件")
        
    try:
        # 验证PDF文件
        validate_pdf_file(file_obj)
        
        # 创建PDF阅读器
        reader = PyPDF2.PdfReader(file_obj)
        
        # 提取文本
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"=== 第{i+1}页 ===\n{page_text}")
            except Exception as e:
                logger.warning(f"无法提取PDF第{i+1}页文本: {e}")
                text_parts.append(f"=== 第{i+1}页 [提取失败] ===\n")
        
        # 合并所有文本
        all_text = "\n\n".join(text_parts)
        
        if not all_text.strip():
            raise DocumentProcessError("PDF文档中未找到有效文本内容")
        
        logger.info(f"成功提取PDF文本，总长度: {len(all_text)} 字符，页数: {len(reader.pages)}")
        return all_text
        
    except DocumentProcessError:
        raise
    except Exception as e:
        logger.error(f"PDF文本提取失败: {e}")
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise DocumentProcessError("PDF文档可能受密码保护，请提供未加密的文档")
        else:
            raise DocumentProcessError(f"PDF文档处理失败: {str(e)}")

def extract_text_from_docx(file_obj) -> str:
    """从DOCX文件中提取文本"""
    try:
        # 验证文件
        validate_docx_file(file_obj)
        
        # 创建Document对象
        doc = Document(file_obj)
        
        # 提取段落文本
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 只保留非空段落
                paragraphs.append(text)
        
        # 提取表格文本
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_texts.append(" | ".join(row_text))
        
        # 合并所有文本
        all_text = "\n".join(paragraphs)
        if table_texts:
            all_text += "\n\n=== 表格内容 ===\n" + "\n".join(table_texts)
        
        if not all_text.strip():
            raise DocumentProcessError("文档中未找到有效文本内容")
        
        logger.info(f"成功提取DOCX文本，总长度: {len(all_text)} 字符，段落数: {len(paragraphs)}, 表格数: {len(doc.tables)}")
        return all_text
        
    except InvalidXmlError as e:
        logger.error(f"DOCX文件XML格式错误: {e}")
        raise DocumentProcessError("文档文件损坏或格式不正确")
    
    except Exception as e:
        logger.error(f"DOCX文本提取失败: {e}")
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise DocumentProcessError("文档可能受密码保护，请提供未加密的文档")
        else:
            raise DocumentProcessError(f"文档处理失败: {str(e)}")

def extract_text(upload_file) -> str:
    """从上传的文件中提取文本，支持DOCX和PDF格式"""
    try:
        # 获取文件扩展名
        file_ext = os.path.splitext(upload_file.filename.lower())[1] if upload_file.filename else ""
        
        # 根据文件类型选择处理方法
        if file_ext == ".pdf":
            return extract_text_from_pdf(upload_file.file)
        elif file_ext == ".docx":
            return extract_text_from_docx(upload_file.file)
        else:
            raise DocumentProcessError(f"不支持的文件类型: {file_ext}，仅支持.docx和.pdf")
            
    except DocumentProcessError:
        raise
    except Exception as e:
        logger.error(f"文本提取失败: {e}")
        raise DocumentProcessError(f"文档处理失败: {str(e)}")

def split_chapters(text: str) -> List[Dict[str, Any]]:
    """将文本按章节分割"""
    if not text or not text.strip():
        return []
    
    chapters = []
    current_chapter = {
        "title": "文档开头",
        "content": "",
        "start_pos": 0,
        "important_score": 0
    }
    
    lines = text.split("\n")
    current_content = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # 检查是否为章节标题
        is_chapter_title = False
        chapter_title = line
        
        for pattern in CHAPTER_PATTERNS:
            match = pattern.match(line)
            if match:
                is_chapter_title = True
                if match.groups():
                    chapter_title = match.group(1).strip() or line
                break
        
        if is_chapter_title and current_content:
            # 保存当前章节
            current_chapter["content"] = "\n".join(current_content)
            current_chapter["important_score"] = calculate_importance_score(current_chapter["content"])
            chapters.append(current_chapter)
            
            # 开始新章节
            current_chapter = {
                "title": chapter_title,
                "content": "",
                "start_pos": i,
                "important_score": 0
            }
            current_content = []
        else:
            current_content.append(line)
    
    # 添加最后一个章节
    if current_content:
        current_chapter["content"] = "\n".join(current_content)
        current_chapter["important_score"] = calculate_importance_score(current_chapter["content"])
        chapters.append(current_chapter)
    
    # 如果没有找到章节，将整个文档作为一个章节
    if not chapters:
        chapters.append({
            "title": "完整文档",
            "content": text,
            "start_pos": 0,
            "important_score": calculate_importance_score(text)
        })
    
    logger.info(f"文档分割完成，共 {len(chapters)} 个章节")
    return chapters

def calculate_importance_score(text: str) -> float:
    """计算文本重要性评分"""
    if not text:
        return 0.0
    
    score = 0.0
    text_lower = text.lower()
    
    # 关键词匹配评分
    for keyword in IMPORTANT_KEYWORDS:
        count = text_lower.count(keyword)
        score += count * 1.0
    
    # 长度评分（较长的章节可能更重要）
    length_score = min(len(text) / 1000, 2.0)  # 最多2分
    score += length_score
    
    # 数字和特殊符号评分（可能包含具体要求）
    number_count = len(re.findall(r'\d+', text))
    score += min(number_count * 0.1, 1.0)  # 最多1分
    
    return round(score, 2)

def extract_key_sections(text: str) -> Dict[str, str]:
    """提取关键章节"""
    chapters = split_chapters(text)
    
    key_sections = {
        "technical_requirements": "",
        "commercial_requirements": "",
        "evaluation_criteria": "",
        "qualification_requirements": "",
        "other_important": ""
    }
    
    # 关键词映射
    section_keywords = {
        "technical_requirements": ["技术要求", "技术规格", "技术参数", "功能要求"],
        "commercial_requirements": ["商务要求", "商务条款", "合同条款", "付款"],
        "evaluation_criteria": ["评分", "评标", "评审", "打分", "加分", "减分"],
        "qualification_requirements": ["资格", "资质", "业绩", "人员", "注册资本"]
    }
    
    for chapter in chapters:
        content = chapter["content"]
        title = chapter["title"].lower()
        
        # 根据标题和内容匹配章节类型
        matched = False
        for section_type, keywords in section_keywords.items():
            if any(keyword in title or keyword in content for keyword in keywords):
                if key_sections[section_type]:
                    key_sections[section_type] += "\n\n" + content
                else:
                    key_sections[section_type] = content
                matched = True
                break
        
        # 如果没有匹配到特定类型，且重要性评分较高，归入其他重要内容
        if not matched and chapter["important_score"] > 3.0:
            if key_sections["other_important"]:
                key_sections["other_important"] += "\n\n" + content
            else:
                key_sections["other_important"] = content
    
    return key_sections

def get_document_stats(text: str) -> Dict[str, Any]:
    """获取文档统计信息"""
    if not text:
        return {}
    
    lines = text.split("\n")
    non_empty_lines = [line for line in lines if line.strip()]
    
    # 计算各种统计信息
    stats = {
        "total_chars": len(text),
        "total_lines": len(lines),
        "non_empty_lines": len(non_empty_lines),
        "avg_line_length": sum(len(line) for line in non_empty_lines) / len(non_empty_lines) if non_empty_lines else 0,
        "keyword_counts": {},
        "chapters_count": len(split_chapters(text))
    }
    
    # 统计关键词出现次数
    text_lower = text.lower()
    for keyword in IMPORTANT_KEYWORDS:
        count = text_lower.count(keyword)
        if count > 0:
            stats["keyword_counts"][keyword] = count
    
    return stats